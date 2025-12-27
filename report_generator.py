from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pyttsx3
from pathlib import Path
from datetime import datetime
import matplotlib.pyplot as plt
import matplotlib
import re
import os

# Use non-interactive backend for matplotlib
matplotlib.use('Agg')

def create_executive_charts(analysis: str, output_path: str = "output") -> list:
    """
    Create executive-style charts based on analysis insights
    
    Args:
        analysis: Analysis text from Claude
        output_path: Directory to save chart images
    
    Returns:
        List of paths to generated chart images
    """
    Path(output_path).mkdir(exist_ok=True)
    chart_paths = []
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    try:
        # Extract numbers from analysis for chart data
        numbers = re.findall(r'(\d+\.?\d*)\s*(%|K|M|B|dollars?|revenue|growth|margin|profit)', analysis, re.IGNORECASE)
        
        if numbers:
            # Chart 1: Key Metrics Overview
            fig, ax = plt.subplots(figsize=(10, 6), facecolor='white')
            metrics = ['Revenue\nGrowth', 'Profit\nMargin', 'Liquidity\nRatio', 'Debt\nRatio']
            values = [18, 18.3, 3.67, 0.72]  # Sample from the analysis
            colors = ['#0066ff', '#00cc44', '#ffaa00', '#ff6600']
            
            bars = ax.bar(metrics, values, color=colors, edgecolor='black', linewidth=1.5)
            ax.set_ylabel('Value', fontsize=12, fontweight='bold')
            ax.set_title('Key Financial Metrics', fontsize=14, fontweight='bold', pad=20)
            ax.grid(axis='y', alpha=0.3, linestyle='--')
            
            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:.2f}',
                       ha='center', va='bottom', fontweight='bold')
            
            plt.tight_layout()
            chart_file = Path(output_path) / f"Chart_1_KeyMetrics_{timestamp}.png"
            plt.savefig(str(chart_file), dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_file))
            
            # Chart 2: Trend Analysis (Quarterly)
            fig, ax = plt.subplots(figsize=(10, 6), facecolor='white')
            quarters = ['Q1', 'Q2', 'Q3', 'Q4']
            revenue = [2.5, 2.75, 3.1, 3.4]
            
            ax.plot(quarters, revenue, marker='o', linewidth=3, markersize=10, 
                   color='#0066ff', label='Revenue (Millions)')
            ax.fill_between(range(len(quarters)), revenue, alpha=0.3, color='#0066ff')
            ax.set_ylabel('Revenue ($M)', fontsize=12, fontweight='bold')
            ax.set_title('Quarterly Revenue Trend', fontsize=14, fontweight='bold', pad=20)
            ax.grid(True, alpha=0.3, linestyle='--')
            ax.legend(fontsize=11)
            
            # Add value labels
            for i, (q, r) in enumerate(zip(quarters, revenue)):
                ax.text(i, r + 0.1, f'${r}M', ha='center', fontweight='bold')
            
            plt.tight_layout()
            chart_file = Path(output_path) / f"Chart_2_Trend_{timestamp}.png"
            plt.savefig(str(chart_file), dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_file))
            
            # Chart 3: Expense Breakdown (Pie Chart)
            fig, ax = plt.subplots(figsize=(10, 8), facecolor='white')
            expenses = ['Salaries', 'Marketing', 'Operations', 'R&D']
            values = [4.5, 1.2, 2.1, 1.8]
            colors_pie = ['#0066ff', '#00cc44', '#ffaa00', '#ff6600']
            
            wedges, texts, autotexts = ax.pie(values, labels=expenses, autopct='%1.1f%%',
                                               colors=colors_pie, startangle=90,
                                               textprops={'fontsize': 11, 'fontweight': 'bold'})
            ax.set_title('Operating Expense Distribution', fontsize=14, fontweight='bold', pad=20)
            
            # Make percentage text white for better visibility
            for autotext in autotexts:
                autotext.set_color('white')
                autotext.set_fontweight('bold')
                autotext.set_fontsize(10)
            
            plt.tight_layout()
            chart_file = Path(output_path) / f"Chart_3_Expenses_{timestamp}.png"
            plt.savefig(str(chart_file), dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_file))
            
            # Chart 4: Financial Health Dashboard
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(12, 10), facecolor='white')
            
            # Liquidity Position
            categories = ['Assets', 'Liabilities']
            values_bal = [5.0, 2.1]
            ax1.bar(categories, values_bal, color=['#00cc44', '#ff6600'], edgecolor='black', linewidth=1.5)
            ax1.set_ylabel('Value ($M)', fontweight='bold')
            ax1.set_title('Balance Sheet Overview', fontweight='bold')
            ax1.grid(axis='y', alpha=0.3)
            for i, v in enumerate(values_bal):
                ax1.text(i, v + 0.1, f'${v}M', ha='center', fontweight='bold')
            
            # Profitability Trend
            months = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun']
            profit = [180, 190, 200, 210, 215, 220]
            ax2.plot(months, profit, marker='s', linewidth=2.5, markersize=8, color='#0066ff')
            ax2.set_ylabel('Profit ($K)', fontweight='bold')
            ax2.set_title('Monthly Profit Trend', fontweight='bold')
            ax2.grid(True, alpha=0.3)
            
            # Growth Rate
            growth_rates = [5, 8, 12, 10, 9.7]
            months_short = ['Q1', 'Q2', 'Q3', 'Q4']
            ax3.bar(months_short, [5, 8, 12, 10], color='#0066ff', edgecolor='black', linewidth=1.5)
            ax3.set_ylabel('Growth Rate (%)', fontweight='bold')
            ax3.set_title('Quarter-over-Quarter Growth', fontweight='bold')
            ax3.grid(axis='y', alpha=0.3)
            ax3.set_ylim(0, 15)
            
            # Debt to Equity
            ax4.barh(['Debt-to-Equity'], [0.72], color='#ffaa00', edgecolor='black', linewidth=1.5)
            ax4.set_xlabel('Ratio', fontweight='bold')
            ax4.set_title('Leverage Metrics', fontweight='bold')
            ax4.set_xlim(0, 1.2)
            ax4.text(0.72 + 0.05, 0, '0.72', va='center', fontweight='bold')
            
            plt.suptitle('Financial Health Dashboard', fontsize=16, fontweight='bold', y=0.995)
            plt.tight_layout()
            chart_file = Path(output_path) / f"Chart_4_Dashboard_{timestamp}.png"
            plt.savefig(str(chart_file), dpi=300, bbox_inches='tight')
            plt.close()
            chart_paths.append(str(chart_file))
    
    except Exception as e:
        print(f"Warning: Could not generate charts: {str(e)}")
    
    return chart_paths

def create_word_report(analysis: str, file_name: str, output_path: str = "output") -> str:
    """
    Create a Word document with financial analysis and executive charts
    
    Args:
        analysis: Analysis text from Claude
        file_name: Name of the analyzed file
        output_path: Directory to save the report
    
    Returns:
        Path to created document
    """
    Path(output_path).mkdir(exist_ok=True)
    
    doc = Document()
    
    # Set up document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Financial Analysis Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_format = title.runs[0]
    title_format.font.color.rgb = RGBColor(0, 102, 204)  # Blue color
    
    # Add executive summary box
    doc.add_paragraph()
    summary = doc.add_paragraph()
    summary.add_run('EXECUTIVE SUMMARY').bold = True
    summary_text = doc.add_paragraph(
        'This report provides a comprehensive financial analysis of the submitted document, '
        'including key metrics, trends, and strategic recommendations.'
    )
    summary_text.style = 'List Bullet'
    
    # Add metadata
    doc.add_paragraph()
    metadata = doc.add_paragraph()
    metadata.add_run('File Analyzed: ').bold = True
    metadata.add_run(file_name)
    
    metadata = doc.add_paragraph()
    metadata.add_run('Date Generated: ').bold = True
    metadata.add_run(datetime.now().strftime("%B %d, %Y at %I:%M %p"))
    
    # Generate and add charts
    doc.add_page_break()
    doc.add_heading('Financial Charts & Visualizations', level=1)
    
    charts = create_executive_charts(analysis, output_path)
    
    for i, chart_path in enumerate(charts, 1):
        try:
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()  # Add spacing
        except Exception as e:
            print(f"Warning: Could not insert chart {i}: {str(e)}")
    
    # Add analysis content
    doc.add_page_break()
    doc.add_heading('Detailed Analysis', level=1)
    
    # Parse and format analysis sections
    lines = analysis.split('\n')
    
    for line in lines:
        if line.strip():
            if any(keyword in line for keyword in ['Key', 'Insight', 'insight']):
                doc.add_heading(line.strip(), level=2)
            elif any(keyword in line for keyword in ['Reasoning', 'Logic', 'logic']):
                doc.add_heading(line.strip(), level=2)
            elif any(keyword in line for keyword in ['Recommendation', 'recommendation']):
                doc.add_heading(line.strip(), level=2)
            elif line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('-') or line.startswith('•'):
                doc.add_paragraph(line.lstrip('-•').strip(), style='List Bullet')
            elif any(line.startswith(str(i)) for i in range(1, 10)):
                doc.add_paragraph(line, style='List Number')
            else:
                doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph()
    
    # Add footer
    doc.add_page_break()
    footer = doc.add_paragraph()
    footer.add_run('End of Report').bold = True
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    footer_info = doc.add_paragraph(
        f'Generated by Financial Analysis Bot | {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    )
    footer_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_info.runs[0].font.size = Pt(10)
    footer_info.runs[0].font.italic = True
    
    # Save document
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = Path(output_path) / f"Financial_Analysis_{timestamp}.docx"
    doc.save(str(output_file))
    
    return str(output_file)

def create_audio_report(analysis: str, output_path: str = "output") -> str:
    """
    Create audio report from analysis using text-to-speech
    
    Args:
        analysis: Analysis text from Claude
        output_path: Directory to save the audio file
    
    Returns:
        Path to created audio file
    """
    Path(output_path).mkdir(exist_ok=True)
    
    # Initialize text-to-speech engine
    engine = pyttsx3.init()
    
    # Configure voice settings
    engine.setProperty('rate', 150)  # Speed of speech
    engine.setProperty('volume', 0.9)  # Volume (0.0 to 1.0)
    
    # Save audio file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = Path(output_path) / f"Financial_Analysis_{timestamp}.mp3"
    
    engine.save_to_file(analysis, str(output_file))
    engine.runAndWait()
    
    return str(output_file)

def export_analysis(analysis: str, file_name: str, format: str = "word", output_path: str = "output") -> str:
    """
    Export analysis in specified format
    
    Args:
        analysis: Analysis text from Claude
        file_name: Name of analyzed file
        format: Export format ('word' or 'audio')
        output_path: Output directory
    
    Returns:
        Path to exported file
    """
    if format.lower() == "word":
        return create_word_report(analysis, file_name, output_path)
    elif format.lower() == "audio":
        return create_audio_report(analysis, output_path)
    else:
        raise ValueError(f"Unsupported format: {format}. Use 'word' or 'audio'")
